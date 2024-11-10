import docx
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from playwright.async_api import async_playwright
from datetime import datetime

from utils.setup_log import logger


def read_docx(file_path):
    """
    This function reads a Word document from the given file path.
    """
    doc = Document(file_path)
    return doc


def get_tables(doc):
    """
    This function extracts all tables from a document and returns them as a list.

    Returns:
    None
    """
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables


def extract_contents_from_tables(tables):
    """
    This function extracts all numbers (inside parentheses) from each row of the tables.

    Returns:
    {company_name: company_code}
    """
    # extracted_numbers = []
    # company_name = []
    contents = {}
    logger.info('[-]尝试读取客户信息...')
    for table in tables:
        # print(table)
        for row in table:
            if row[1]:
                try:
                    numbers = re.findall(r'（(\d+)）', row[1])[0]# Find numbers inside parentheses
                    company_name = re.findall(r'([^\（]+)（', row[1])[0]
                except Exception as e:
                    logger.info(e)
                    numbers = re.findall(r'\((\d+)\)', row[1])[0]  # Find numbers inside parentheses
                    company_name = re.findall(r'([^\(]+)\(', row[1])[0]
                contents[company_name] = numbers
    logger.success('[+]成功读取客户信息！')
    return contents


async def add_hyperlink(paragraph, url, text):
    """
    将超链接添加到 Word 文档的段落中，并应用蓝色和下划线样式。

    Args:
    paragraph (docx.text.Paragraph): 要添加超链接的段落。
    url (str): 超链接的目标 URL。
    text (str): 显示的文本内容。

    Returns:
    None
    """
    # 创建关系
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # 创建超链接元素
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # 创建显示文本的 run
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 设置字体颜色为蓝色
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # 蓝色的十六进制表示
    rPr.append(color)

    # 设置下划线
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')  # 单线下划线
    rPr.append(underline)
    run.append(rPr)

    # 设置字体为宋体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '宋体')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)

    # 设置字号为小四号（对应12号字体，Word中的值为32点）
    size = OxmlElement('w:sz')
    size.set(qn('w:val'), '32')  # 小四号字体的值为32
    rPr.append(size)

    # 添加超链接的显示文本
    run_text = OxmlElement('w:t')
    run_text.text = text
    run.append(run_text)

    hyperlink.append(run)
    paragraph._element.append(hyperlink)  # 将超链接添加到段落中


class Doc:
    def __init__(self, cninfo_hyperlinks, sina_hyperlinks, saved_path):
        """
        Args:
        hyperlinks (dict): {company_name: {title: {href: date}} }
        saved_path (str): The path to save the output .docx file.
        """
        self.cn_hyperlinks = cninfo_hyperlinks
        self.sn_hyperlinks = sina_hyperlinks
        self.saved_path = saved_path
        self.d = Document()  # Create a new Document

    async def add_beginnings_to_docx(self):
        """
        Create a new document.

        Returns:
        None
        """
        # Step 1: Add formatted beginnings
        date_today = datetime.today().strftime('%Y年%m月%d日 %H:%M')
        self.d.add_paragraph(f"各位律师好，\n{date_today}，常年客户信息披露更新如下：\n")

    # Function to write extracted links into a formatted Word document
    async def write_announce_to_docx(self):
        """
        Writes a list of extracted hrefs (URLs) into a Word document with a formatted title and beginning text.

        Returns:
        None
        """

        self.d.add_heading('一、公告：', level=2)
        logger.info('[-]尝试写入更新的公告...')
        # Iterate over dictionary items
        for index, (company_name, contents) in enumerate(self.cn_hyperlinks.items(), start=1):
            # Write the company name
            self.d.add_paragraph(f'{index}. {company_name}')  # Write company name

            # Iterate over titles and dates within the company data
            for title, dates in contents.items():
                for date, url in dates.items():
                    # Add a new paragraph for the date and title
                    nl = self.d.add_paragraph(f"{date}:")  # Add date

                    # Add the hyperlink asynchronously
                    await add_hyperlink(nl, url, title)  # Add the hyperlink with title text

        self.d.save(self.saved_path)

    # Function to write extracted links into a formatted Word document
    async def write_news_to_docx(self):
        """
        Writes a list of extracted hrefs (URLs) into a Word document with a formatted title and beginning text.

        Returns:
        None
        """
        # d = Document(self.saved_path)
        self.d.add_heading('二、新闻：', level=2)
        logger.info('[-]尝试写入更新的公告...')
        for index, (name, contents) in enumerate(self.sn_hyperlinks.items(), start=1):
            if contents:
                self.d.add_paragraph(f'{index}.' + name)  # Head: 1. Company Name
                for title, url in contents.items():
                    nl = self.d.add_paragraph()
                    await add_hyperlink(nl, url, title)

        logger.success('[+]公告成功写入！')
        self.d.save(self.saved_path)
        logger.info(f"Document saved to {self.saved_path}")

    async def change_font(self):
        logger.info("[-]尝试改变字体...")
        for paragraph in self.d.paragraphs:
            for run in paragraph.runs:
                run.font.name = '宋体'

                # 处理中文字体名的兼容性问题
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                # 设置小四号字体，对应12pt
                run.font.size = Pt(16)
        logger.success("[+]成功改变字体！")
        self.d.save(self.saved_path)


